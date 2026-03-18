"""Invoice generation service - Word document (Delphinus format)."""
import logging
import subprocess
import tempfile
import os
from datetime import datetime
from io import BytesIO
from typing import Optional, Dict, Any

logger = logging.getLogger(__name__)

from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.utils.number_to_words import number_to_french_words


def generate_invoice_word(
    document,
    awb_details: Optional[Dict[str, Any]],
    amount_usd: float,
    usd_to_gnf: int,
) -> bytes:
    """
    Generate invoice as Word document (Delphinus format).
    Returns docx file as bytes.
    """
    doc = DocxDocument()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    def add_para(text: str, size_pt: int = 12, bold: bool = False, align: str = None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size_pt)
        run.bold = bold
        if align == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        return p
    
    # Date
    doc_date = document.document_date.strftime('%d %B %Y') if document.document_date else '-'
    # French month names
    months_fr = {'January': 'janvier', 'February': 'février', 'March': 'mars', 'April': 'avril',
                 'May': 'mai', 'June': 'juin', 'July': 'juillet', 'August': 'août',
                 'September': 'septembre', 'October': 'octobre', 'November': 'novembre',
                 'December': 'décembre'}
    for en, fr in months_fr.items():
        doc_date = doc_date.replace(en, fr)
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_date = date_para.add_run(f'Conakry, le {doc_date}')
    run_date.font.name = 'Times New Roman'
    run_date.font.size = Pt(14)
    
    ref = document.reference_number or ''
    date_part = datetime.now().strftime('%d/%m/%y')
    invoice_number = f"{ref}/EC/JA/{date_part}" if ref else '-'
    p_inv = doc.add_paragraph()
    run_inv = p_inv.add_run(f"Facture N° {invoice_number}")
    run_inv.font.name = 'Times New Roman'
    run_inv.font.size = Pt(14)
    run_inv.bold = True
    run_inv.underline = True
    doc.add_paragraph()
    
    # Client (right aligned, bold, same line: Client : [name])
    client_name = document.consignee or document.shipper or 'Client'
    client_para = doc.add_paragraph()
    client_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = client_para.add_run(f'Client : {client_name}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    if (awb_details or {}).get('consignee_details'):
        run2 = client_para.add_run(f'\n{awb_details["consignee_details"]}')
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
    doc.add_paragraph()
    
    # Route
    awb_details = awb_details or {}
    origin_code = (awb_details.get('routing', {}).get('departure_code') or document.origin or '---')[:3].upper()
    to_list = awb_details.get('routing', {}).get('to') or []
    dest_code = (to_list[-1] if to_list else document.destination or '---')[:3].upper()
    route = f'{origin_code}-{dest_code}'.replace('---', '-')
    
    total_pieces = 0
    total_weight = 0.0
    weight_scale = 'K'
    nature_desc = 'Transport aérien'
    if awb_details:
        rd = awb_details.get('rate_description') or {}
        total_pieces = rd.get('total_pieces') or awb_details.get('total_pieces') or 0
        total_weight = rd.get('total_weight') or awb_details.get('total_weight') or 0.0
        weight_scale = rd.get('weight_scale') or awb_details.get('weight_scale') or 'K'
        items = rd.get('items') or []
        if items:
            nature_desc = items[0].get('nature') or 'Marchandises'
    weight_label = 'Kg' if weight_scale == 'K' else weight_scale
    
    pieces_part = f'{total_pieces} colis' if total_pieces else ''
    weight_part = f'{total_weight} {weight_label}' if total_weight else ''
    libelle = f"Transport de {pieces_part} {weight_part} de {nature_desc} {route}".strip()
    
    nature_para = doc.add_paragraph()
    nature_para.paragraph_format.space_before = Pt(0)
    nature_para.paragraph_format.space_after = Pt(0)
    run_nat1 = nature_para.add_run("Nature de l'opération")
    run_nat1.font.name = 'Times New Roman'
    run_nat1.font.size = Pt(12)
    run_nat1.bold = True
    run_nat1.underline = True
    run_nat2 = nature_para.add_run(f" : {libelle}")
    run_nat2.font.name = 'Times New Roman'
    run_nat2.font.size = Pt(12)
    run_nat2.bold = True
    lta_para = doc.add_paragraph()
    lta_para.paragraph_format.space_before = Pt(0)
    lta_para.paragraph_format.space_after = Pt(0)
    run_lta = lta_para.add_run(f"LTA : {document.document_number or '-'}")
    run_lta.font.name = 'Times New Roman'
    run_lta.font.size = Pt(12)
    run_lta.italic = True
    doc.add_paragraph()
    doc.add_paragraph()
    
    montant_para = doc.add_paragraph()
    montant_para.paragraph_format.space_before = Pt(0)
    montant_para.paragraph_format.space_after = Pt(0)
    run_m1 = montant_para.add_run("Montant Total de l'opération")
    run_m1.font.name = 'Times New Roman'
    run_m1.font.size = Pt(12)
    run_m1.bold = True
    run_m1.underline = True
    run_m2 = montant_para.add_run(f" : {amount_usd:,.2f} USD (1 USD = {usd_to_gnf:,} GNF)")
    run_m2.font.name = 'Times New Roman'
    run_m2.font.size = Pt(12)
    doc.add_paragraph()
    
    # Table
    currency = (awb_details or {}).get('currency') or 'USD'
    total_gnf = round(amount_usd * usd_to_gnf)
    
    table = doc.add_table(rows=3, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    headers = ['N°', 'LIBELLE', 'NOMBRE', f'Montant en {currency}', 'MONTANT en GNF']
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
    
    data_row = table.rows[1].cells
    data_row[0].text = '1'
    data_row[1].text = f"{libelle} {document.document_number or ''}"
    data_row[2].text = f'{total_pieces} colis'
    data_row[3].text = f'{amount_usd:,.2f}'
    data_row[4].text = f'{total_gnf:,}'
    for cell in data_row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
    
    total_row = table.rows[2].cells
    total_row[0].merge(total_row[1])
    total_row[0].merge(total_row[2])
    total_row[0].merge(total_row[3])
    total_row[0].text = 'MONTANT TOTAL'
    total_row[4].text = f'{total_gnf:,}'
    for cell in total_row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Montant total à Payer (bold, 14pt)
    p_payer = doc.add_paragraph()
    run_payer = p_payer.add_run(f"Montant total à Payer : ………………………. {total_gnf:,} GNF")
    run_payer.font.name = 'Times New Roman'
    run_payer.font.size = Pt(14)
    run_payer.bold = True
    doc.add_paragraph()
    
    # Amount in words (phrase size 14, amount part in italic + bold)
    amount_words = number_to_french_words(total_gnf) + ' francs guinéens.'
    p = doc.add_paragraph()
    run1 = p.add_run("Arrêtée la présente facture à la somme de : ")
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(14)
    run2 = p.add_run(amount_words)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(14)
    run2.italic = True
    run2.bold = True
    doc.add_paragraph()
    
    # Signature (bold, 14pt)
    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sig.add_run('Le Service Administratif & Financier')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _convert_docx_to_pdf_libreoffice(docx_path: str, output_dir: str) -> str:
    """Convert docx to PDF using LibreOffice (soffice). Returns path to PDF."""
    for cmd in ["libreoffice", "soffice"]:
        try:
            result = subprocess.run(
                [cmd, "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode == 0:
                pdf_path = os.path.join(output_dir, "facture.pdf")
                if os.path.exists(pdf_path):
                    return pdf_path
        except FileNotFoundError:
            continue
        except subprocess.TimeoutExpired:
            raise RuntimeError(f"Conversion timeout. LibreOffice ({cmd}) took too long.")
    raise RuntimeError(
        "LibreOffice non trouvé. Installez-le : apt-get install libreoffice (Linux) "
        "ou brew install libreoffice (Mac)"
    )


def generate_invoice_pdf(
    document,
    awb_details: Optional[Dict[str, Any]],
    amount_usd: float,
    usd_to_gnf: int,
) -> bytes:
    """
    Generate invoice as PDF (Word converted to PDF for identical format).
    Returns pdf file as bytes.
    Requires LibreOffice (Linux/Mac) or Microsoft Word (Windows).
    """
    docx_bytes = generate_invoice_word(document, awb_details, amount_usd, usd_to_gnf)

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "facture.docx")
        pdf_path = os.path.join(tmpdir, "facture.pdf")

        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        converted = False
        # Try docx2pdf first (uses Word on Windows, LibreOffice on Mac)
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            converted = True
        except Exception as e:
            logger.warning("docx2pdf failed: %s, trying LibreOffice directly", e)

        # Fallback: LibreOffice directly (Linux, Mac)
        if not converted or not os.path.exists(pdf_path):
            try:
                pdf_path = _convert_docx_to_pdf_libreoffice(docx_path, tmpdir)
            except RuntimeError as e:
                raise RuntimeError(str(e))

        if not os.path.exists(pdf_path):
            pdf_path = os.path.join(tmpdir, "facture.pdf")

        with open(pdf_path, "rb") as f:
            return f.read()
